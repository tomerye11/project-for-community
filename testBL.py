from email.mime.application import MIMEApplication
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx2pdf import convert
from datetime import datetime
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

app = Flask(__name__)
CORS(app)  # Add this line to enable CORS for all routes

def replace_text_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for key, value in replacements.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key, value)

def replace_text_in_docx(arr, docx_path, output_path):
    doc = Document(docx_path)
    replacements = {
        'Test1': arr[0],
        'Test2': arr[1],
        'Test3': arr[2],
        'Test4': arr[3],
        'Test5': arr[4],
        'Test6': arr[5],
        'date': datetime.now().strftime('%d/%m/%Y')
    }
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        replace_text_in_table(table, replacements)
    temp_docx_path = output_path.replace('.pdf', '_temp.docx')
    doc.save(temp_docx_path)
    convert(temp_docx_path, output_path)
    os.remove(temp_docx_path)

@app.route('/generate_pdf', methods=['POST'])
def generate_pdf():
    data = request.json
    arr = data['arr']
    directory = 'C:\\BituahLeumiForms'

    if not os.path.exists(directory):
        os.makedirs(directory)

    output_path = os.path.join(directory, f'{arr[2]}.pdf')
    replace_text_in_docx(arr, 'volDoc.docx', output_path)
    
    return send_file(output_path, as_attachment=True, download_name=f'{arr[2]}.pdf')

def send_email(to_email, subject, message, attachment_path=None):
    from_email = 'minhalTester@gmail.com'
    password = 'luxf scio xffd ldtc'

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject

    msg.attach(MIMEText(message, 'html'))

    if attachment_path:
        with open(attachment_path, 'rb') as attachment:
            part = MIMEApplication(attachment.read(), Name=os.path.basename(attachment_path))
        part['Content-Disposition'] = f'attachment; filename="{os.path.basename(attachment_path)}"'
        msg.attach(part)

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        text = msg.as_string()
        server.sendmail(from_email, to_email, text)
        server.quit()
        print("Email sent successfully!")
    except Exception as e:
        print(f"Failed to send email: {e}")

@app.route('/approve_volunteer', methods=['POST'])
def approve_volunteer():
    data = request.get_json()
    volunteer_email = data.get('email')
    pdf_path = data.get('pdf_path')
    whatsAppLink = data.get('whatsAppLink')
    
    html_message = f'''
    <!DOCTYPE html>
    <html lang="he">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>אישור התנדבות</title>
    </head>
    <body>
        <p>שלום רב,</p>
        <p>
            ברצוננו להודיעך כי אושרת כמתנדב עבור המינהל הקהילתי לב העיר!
        </p>
        <p>
            לינק לקבוצת הוואצאפ של התחום שבחרת:
        </p>
        <p>
            <a href="{whatsAppLink}">לחץ כאן להצטרפות לקבוצת הוואצאפ</a>
        </p>
        <p>
            מצורף קובץ הביטוח הלאומי שמולא עבורך.
        </p>
        <p>
            מאחלים בהצלחה ושמחים שהצטרפת אלינו!
        </p>
    </body>
    </html>
    '''

    send_email(volunteer_email, 'ברוך הבא! אושרת כמתנדב', html_message, attachment_path=pdf_path)

    return jsonify({'status': 'success', 'message': 'Volunteer approved and email sent'})


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
